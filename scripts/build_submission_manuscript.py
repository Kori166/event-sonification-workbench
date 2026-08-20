"""Build the private, submission-formatted dissertation DOCX from Phase D sources."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
CHAPTERS = REPO / "docs" / "dissertation" / "chapters"
FULL_TITLE = "A Reproducible Workbench for Event-Based Sonification of Annotated Video Datasets"
FIGURE_PNGS = {
    "../figures/figure-1-architecture-and-provenance.svg": "figure-1-architecture-and-provenance.png",
    "../../evaluation/reporting/figures/figure-1-event-outcomes.svg": "figure-1-event-outcomes.png",
    "../../evaluation/reporting/figures/figure-2-cue-density.svg": "figure-2-cue-density.png",
    "../../evaluation/reporting/figures/figure-3-overlap-burden.svg": "figure-3-overlap-burden.png",
}
TABLE_LINKS = {
    "../../evaluation/reporting/tables/table-1-event-accounting-and-coverage.md": REPO / "docs" / "evaluation" / "reporting" / "tables" / "table-1-event-accounting-and-coverage.md",
    "../../evaluation/reporting/tables/table-2-timing-traceability-reproducibility.md": REPO / "docs" / "evaluation" / "reporting" / "tables" / "table-2-timing-traceability-reproducibility.md",
    "../../evaluation/reporting/tables/table-3-density-and-overlap.md": REPO / "docs" / "evaluation" / "reporting" / "tables" / "table-3-density-and-overlap.md",
}
WORD_RE = re.compile(r"\b[\w][\w'’\-]*\b", re.UNICODE)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    tr_pr.append(cannot_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    repeat_header(row)


def set_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run._r.append(node)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("*", "").replace("`", "")


TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^]]+\]\([^)]+\))")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            label = re.match(r"\[([^]]+)\]", token).group(1)
            paragraph.add_run(strip_markdown(label))
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(paragraph, strip_markdown(text))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]], wide: bool = False) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Dissertation Table"
    table.autofit = False
    set_fixed_table_layout(table)
    if cols == 4:
        widths = [Inches(1.92), Inches(1.48), Inches(1.39), Inches(1.39)] if wide else [Inches(1.75), Inches(1.35), Inches(1.54), Inches(1.54)]
    elif cols == 3:
        widths = [Inches(1.35), Inches(2.25), Inches(2.58)]
    elif cols == 2:
        widths = [Inches(1.42), Inches(4.76)]
    else:
        widths = [Inches(6.18 / cols)] * cols
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(7.4 if wide or len(rows) > 8 else 8.2)
                if r_idx == 0:
                    run.bold = True


def add_external_table(doc: Document, path: Path) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    table_start = next(i for i, line in enumerate(lines) if line.lstrip().startswith("|"))
    rows, after = parse_table(lines, table_start)
    add_table(doc, rows, wide=True)
    notes = [line.strip() for line in lines[after:] if line.strip()]
    if notes:
        note = doc.add_paragraph(style="Table Note")
        add_inline(note, "Note. " + strip_markdown(" ".join(notes)))


def set_image_alt(shape, title: str, description: str) -> None:
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc: Document, png_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width = Inches(4.9 if "architecture-and-provenance" in png_path.name else 6.12)
    shape = run.add_picture(str(png_path), width=width)
    title = caption.split(".", 1)[0]
    set_image_alt(shape, title, caption)
    paragraph.paragraph_format.keep_with_next = True
    add_caption(doc, caption)


def collect_navigation() -> tuple[list[tuple[int, str]], list[str], list[str]]:
    def concise_label(label: str) -> str:
        return re.split(
            r" (?=Author-created|Author summary|Contract `|Fixed baseline|Not a listener|Presentation derivative|Sample/seconds|Frozen preset)",
            label,
            maxsplit=1,
        )[0]

    headings: list[tuple[int, str]] = [(1, "Abstract")]
    figures: list[str] = []
    tables: list[str] = []
    for index in range(1, 10):
        path = next(CHAPTERS.glob(f"{index:02d}-*.md"))
        for line in path.read_text(encoding="utf-8").splitlines():
            if line.startswith("# "):
                headings.append((1, strip_markdown(line[2:])))
            elif line.startswith("## ") and index < 9:
                headings.append((2, strip_markdown(line[3:])))
            elif line.startswith("!["):
                figures.append(concise_label(line[2:line.index("](")]))
            elif re.match(r"\*\*Table \d+\.", line) or re.match(r"\[\*\*Table \d+\.", line):
                label = line
                if label.startswith("["):
                    label = label[1:label.index("](")]
                tables.append(concise_label(strip_markdown(label)))
    return headings, figures, tables


def extract_page_map(pdf_path: Path) -> dict[str, int]:
    from pypdf import PdfReader

    headings, figures, tables = collect_navigation()
    labels = [label for _, label in headings] + figures + tables
    pages = PdfReader(str(pdf_path)).pages

    def normalise(value: str) -> str:
        value = unicodedata.normalize("NFKD", value).casefold()
        return "".join(char for char in value if char.isalnum())

    page_text = [normalise(page.extract_text() or "") for page in pages]
    result: dict[str, int] = {}
    for label in labels:
        if label == "Abstract":
            result[label] = 2
            continue
        needle = normalise(label)
        probes = (needle, needle[:100], needle[:70], needle[:45])
        for page_number, text in enumerate(page_text[5:], 6):
            if any(len(probe) >= 8 and probe in text for probe in probes):
                result[label] = page_number
                break
    missing = [label for label in labels if label not in result]
    if missing:
        raise RuntimeError("Could not locate navigation labels in rendered PDF: " + "; ".join(missing))
    return result


def add_static_nav(doc: Document, title: str, entries: Iterable[tuple[int, str]], page_map: dict[str, int]) -> None:
    heading = doc.add_paragraph(title, style="Front Matter Heading")
    heading.paragraph_format.space_after = Pt(12)
    for level, label in entries:
        if title == "Contents":
            style_name = "TOC 1" if level == 1 else "TOC 2"
        else:
            style_name = "Front Matter Entry"
        paragraph = doc.add_paragraph(style=style_name)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.12), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_inline(paragraph, label)
        paragraph.add_run("\t" + (str(page_map[label]) if label in page_map else "—"))


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (("Heading 1", 16, 0, 10), ("Heading 2", 12.5, 12, 5), ("Heading 3", 11, 10, 4)):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    for name in ("TOC 1", "TOC 2"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = Pt(2)
    styles["TOC 1"].font.bold = True
    styles["TOC 2"].paragraph_format.left_indent = Cm(0.65)

    if "Front Matter Entry" not in styles:
        styles.add_style("Front Matter Entry", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Front Matter Entry"]
    style.font.name = "Aptos"
    style.font.size = Pt(9)
    style.font.bold = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.05
    style.paragraph_format.space_after = Pt(4)

    if "Front Matter Heading" not in styles:
        style = styles.add_style("Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Front Matter Heading"]
    style.font.name = "Aptos Display"
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.keep_with_next = True

    if "Dissertation Table" not in styles:
        style = styles.add_style("Dissertation Table", WD_STYLE_TYPE.TABLE)
    styles["Dissertation Table"].base_style = styles["Table Grid"]

    if "Table Note" not in styles:
        style = styles.add_style("Table Note", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Table Note"]
    style.font.name = "Aptos"
    style.font.size = Pt(7.5)
    style.font.italic = True
    style.paragraph_format.space_after = Pt(7)

    if "Reference" not in styles:
        styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Reference"]
    style.base_style = styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(9.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.left_indent = Cm(0.65)
    style.paragraph_format.first_line_indent = Cm(-0.65)
    style.paragraph_format.line_spacing = 1.05
    style.paragraph_format.space_after = Pt(5)

    list_number = styles["List Number"]
    list_number.font.name = "Aptos"
    list_number.font.size = Pt(10.5)
    list_number.paragraph_format.left_indent = Cm(0.75)
    list_number.paragraph_format.first_line_indent = Cm(-0.5)
    list_number.paragraph_format.space_after = Pt(3)


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Event-Based Sonification Workbench")
    hr.font.name = "Aptos"
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(96, 96, 96)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(FULL_TITLE)
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("An MSc Data Science dissertation")
    sr.font.name = "Aptos"
    sr.font.size = Pt(13)
    sr.italic = True
    doc.add_paragraph()
    for text, bold in (("Kori Flowers", True), ("Student ID: 24046378", False), ("MSc Data Science", False), ("UFCF9Y-60-M CSCT Masters Project", False), ("University of the West of England", False), ("August 2026", False)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Aptos"
        r.font.size = Pt(11)
        r.bold = bold
    doc.add_paragraph()
    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = statement.add_run("A dissertation submitted in partial fulfilment of the requirements for the degree of MSc Data Science")
    st.font.size = Pt(9.5)
    st.italic = True
    count = doc.add_paragraph()
    count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count.add_run("Assessed body word count: pending")
    add_bookmark(count, "assessed_word_count", 1)
    doc.add_page_break()


def add_markdown_content(doc: Document, path: Path, asset_dir: Path, skip_first_heading: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    first_heading_skipped = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            if skip_first_heading and not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline(paragraph, line[2:])
        elif line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[3:])
        elif line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[4:])
        elif line.startswith("!["):
            match = re.match(r"!\[(.+)\]\((.+)\)", line)
            caption, target = match.groups()
            add_figure(doc, asset_dir / FIGURE_PNGS[target], caption)
        elif re.match(r"\[\*\*Table \d+\.", line):
            label, target = re.match(r"\[(.+)\]\((.+)\)", line).groups()
            add_caption(doc, strip_markdown(label))
            add_external_table(doc, TABLE_LINKS[target])
        elif re.match(r"\*\*Table \d+\.", line):
            add_caption(doc, strip_markdown(line))
        elif line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        elif re.match(r"\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\. ", "", line))
        else:
            paragraph = doc.add_paragraph(style="Reference" if path.name.startswith("09-") else None)
            add_inline(paragraph, line)
        i += 1


def count_and_stamp_assessed_body(doc: Document) -> int:
    body = doc._element.body
    in_body = False
    words: list[str] = []
    for child in body.iterchildren():
        text = " ".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if text == "1. Introduction":
            in_body = True
        if text == "References":
            in_body = False
        if in_body:
            words.extend(WORD_RE.findall(text))
    count = len(words)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Assessed body word count:"):
            paragraph.clear()
            paragraph.add_run(f"Assessed body word count: {count:,}")
            break
    return count


def build(output: Path, asset_dir: Path, page_map: dict[str, int]) -> int:
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    props = doc.core_properties
    props.title = FULL_TITLE
    props.subject = "MSc Data Science dissertation"
    props.author = "Kori Flowers"
    props.keywords = "sonification, annotated video, reproducibility, provenance, technical evaluation"
    props.comments = "Submission manuscript assembled from the versioned Phase D dissertation sources."

    add_title_page(doc)
    abstract_heading = doc.add_paragraph("Abstract", style="Front Matter Heading")
    abstract_heading.paragraph_format.space_after = Pt(10)
    add_markdown_content(doc, CHAPTERS / "00-abstract.md", asset_dir, skip_first_heading=True)
    doc.add_page_break()

    headings, figures, tables = collect_navigation()
    add_static_nav(doc, "Contents", headings, page_map)
    doc.add_page_break()
    add_static_nav(doc, "List of Figures", [(1, item) for item in figures], page_map)
    doc.add_page_break()
    add_static_nav(doc, "List of Tables", [(1, item) for item in tables], page_map)
    doc.add_page_break()

    for index in range(1, 10):
        path = next(CHAPTERS.glob(f"{index:02d}-*.md"))
        add_markdown_content(doc, path, asset_dir)

    count = count_and_stamp_assessed_body(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    metadata = {
        "title": props.title,
        "author": props.author,
        "output": str(output),
        "assessed_body_word_count": count,
        "assessed_scope": "Chapters 1–8, including headings, captions and table content; excluding title/front matter, abstract and references",
        "page_map": page_map,
    }
    output.with_suffix(".build.json").write_text(json.dumps(metadata, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    return count


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--asset-dir", type=Path, required=True)
    parser.add_argument("--page-map", type=Path)
    parser.add_argument("--extract-page-map-from", type=Path)
    args = parser.parse_args()
    if args.page_map and args.extract_page_map_from:
        parser.error("choose either --page-map or --extract-page-map-from")
    if args.extract_page_map_from:
        page_map = extract_page_map(args.extract_page_map_from.resolve())
    else:
        page_map = json.loads(args.page_map.read_text(encoding="utf-8")) if args.page_map else {}
    count = build(args.output.resolve(), args.asset_dir.resolve(), page_map)
    print(json.dumps({"output": str(args.output.resolve()), "assessed_body_word_count": count}, ensure_ascii=False))


if __name__ == "__main__":
    main()
